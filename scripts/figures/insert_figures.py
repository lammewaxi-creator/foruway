#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Insert figures into the Word document at correct positions."""
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')
doc = Document('小论文5.14_revised_v5.docx')

FIGURES = [
    {
        'path': 'results/journal/figures_new/Fig1_Framework.png',
        'caption': '图 1  ρ感知自适应调度框架',
        'insert_after_text': '3. 研究方法',  # Insert after this heading
    },
    {
        'path': 'results/journal/figures_new/Fig3_Algorithm_Comparison.png',
        'caption': '图 3  算法性能对比（R2 / 1120 任务 / decoupled 模式）',
        'insert_after_text': '注：W 改善率 = (W_static',
    },
    {
        'path': 'results/journal/figures_new/Fig4_Pymoo_Comparison.png',
        'caption': '图 4  与现代多目标进化算法对比（R2 / 600 任务 / pop=30 / gen=30）',
        'insert_after_text': '而是面向 RGV-Lift 耦合调度的专用进化算子',
    },
    {
        'path': 'results/journal/figures_new/Fig6_Hyperparam_Robustness.png',
        'caption': '图 6  超参数稳健性扫描（R2 / 1120 任务 / sync / 3 seeds）',
        'insert_after_text': '注：本表数据系 3 种子',
    },
    {
        'path': 'results/journal/figures_new/Fig5_Ablation_Waterfall.png',
        'caption': '图 5  消融实验模块贡献分析（R2 / 1120 任务 / decoupled / 5 seeds）',
        'insert_after_text': '注：所有数据均为 5 次独立运行实测值',
    },
    {
        'path': 'results/journal/figures_new/Fig7_Lift_Sensitivity.png',
        'caption': '图 7  提升机数量灵敏度分析（R2 / 1120 任务 / decoupled / 3 seeds）',
        'insert_after_text': '注：lift_count = 5 在原始实验设计中被跳过',
    },
]

def set_chinese_font(run, font_name='宋体', size=Pt(10.5)):
    run.font.name = font_name
    run.font.size = size
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def insert_figure_after(doc, target_text, img_path, caption, max_search=300):
    """Find paragraph containing target_text and insert figure after it."""
    for i, para in enumerate(doc.paragraphs):
        if target_text in para.text:
            print(f'  Found target at Para {i}: "{para.text[:60]}..."')
            # Insert a new paragraph after this one
            new_p = para._element.addnext(doc.add_paragraph()._element)
            # Get the paragraph object for the new element
            # Need to find it in doc.paragraphs
            # Actually, addnext inserts the element, but we need to add the picture
            # Let's use a different approach
            return i
    print(f'  WARNING: Could not find text: "{target_text[:40]}..."')
    return None

# Use a simpler approach: insert at specific paragraph indices
# First, let's identify exact paragraph indices

print("Scanning document for insertion points...")
insert_points = {}
for fig in FIGURES:
    target = fig['insert_after_text']
    for i, para in enumerate(doc.paragraphs):
        if target in para.text:
            insert_points[fig['path']] = i
            print(f"  {fig['caption']}: after Para {i}")
            break
    else:
        print(f"  NOT FOUND: {fig['caption']}")

# Now insert figures using XML manipulation
from docx.oxml import OxmlElement

for fig in FIGURES:
    path = fig['path']
    caption = fig['caption']
    target = fig['insert_after_text']

    # Find insertion point
    insert_idx = None
    for i, para in enumerate(doc.paragraphs):
        if target in para.text:
            insert_idx = i
            break

    if insert_idx is None:
        print(f"SKIP: Could not find insertion point for {caption}")
        continue

    # Get the XML element of the target paragraph
    target_elem = doc.paragraphs[insert_idx]._element

    # Create caption paragraph
    cap_p = OxmlElement('w:p')
    cap_r = OxmlElement('w:r')
    cap_rPr = OxmlElement('w:rPr')
    cap_rFonts = OxmlElement('w:rFonts')
    cap_rFonts.set(qn('w:eastAsia'), '宋体')
    cap_rFonts.set(qn('w:ascii'), 'Times New Roman')
    cap_rPr.append(cap_rFonts)
    cap_sz = OxmlElement('w:sz')
    cap_sz.set(qn('w:val'), '21')  # 10.5pt
    cap_rPr.append(cap_sz)
    cap_r.append(cap_rPr)
    cap_t = OxmlElement('w:t')
    cap_t.set(qn('xml:space'), 'preserve')
    cap_t.text = caption
    cap_r.append(cap_t)
    cap_p.append(cap_r)

    # Create figure placeholder paragraph (we'll add image via docx API)
    fig_p = OxmlElement('w:p')
    fig_pPr = OxmlElement('w:pPr')
    fig_jc = OxmlElement('w:jc')
    fig_jc.set(qn('w:val'), 'center')
    fig_pPr.append(fig_jc)
    fig_p.append(fig_pPr)

    # Insert figure paragraph and caption after target
    target_elem.addnext(fig_p)
    fig_p.addnext(cap_p)

    # Now find the figure paragraph in doc.paragraphs and add the image
    # We need to find it by matching the element
    for para in doc.paragraphs:
        if para._element is fig_p:
            run = para.add_run()
            run.add_picture(path, width=Cm(14))
            break

    print(f"INSERTED: {caption}")

# Save
doc.save('小论文5.14_revised_v6.docx')
print('\nSaved to 小论文5.14_revised_v6.docx')
